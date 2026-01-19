from docx import Document
import json

doc = Document('HRM.docx')

data = {
    'paragraphs': [],
    'tables': []
}

for para in doc.paragraphs:
    if para.text.strip():
        data['paragraphs'].append(para.text)

for i, table in enumerate(doc.tables):
    table_data = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        table_data.append(cells)
    data['tables'].append(table_data)

# Write to JSON file
with open('hrm_content.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, indent=2)

# Write to text file
with open('hrm_content.txt', 'w', encoding='utf-8') as f:
    f.write("=" * 80 + "\n")
    f.write("PARAGRAPHS:\n")
    f.write("=" * 80 + "\n")
    for para in data['paragraphs']:
        f.write(para + "\n")
    
    f.write("\n" + "=" * 80 + "\n")
    f.write("TABLES:\n")
    f.write("=" * 80 + "\n")
    for i, table in enumerate(data['tables'], 1):
        f.write(f"\n--- Table {i} ---\n")
        for row in table:
            f.write(" | ".join(row) + "\n")

print("Content exported to hrm_content.json and hrm_content.txt")
